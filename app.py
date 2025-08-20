from flask import Flask, request, render_template, send_file, redirect, url_for
from docx import Document
import os
import random
from werkzeug.utils import secure_filename

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated'
ALLOWED_EXTENSIONS = {'docx'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_questions_from_all_files():
    questions = []
    for filename in os.listdir(UPLOAD_FOLDER):
        if allowed_file(filename):
            doc = Document(os.path.join(UPLOAD_FOLDER, filename))
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 6:
                        row_data = [cell.text.strip() for cell in row.cells[:6]]
                        questions.append(row_data)
    return questions

def generate_docx(questions, output_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=6)
    for row in questions:
        cells = table.add_row().cells
        for i in range(6):
            cells[i].text = row[i]
    doc.save(output_path)

@app.route('/')
def home():
    return redirect(url_for('user_request'))

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        file = request.files.get('file')
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(UPLOAD_FOLDER, filename))
            return "تم رفع الملف بنجاح!"
        return "صيغة غير مدعومة"
    return render_template("admin_upload.html")

@app.route('/generate', methods=['GET', 'POST'])
def user_request():
    if request.method == 'POST':
        count = int(request.form.get('count', 5))
        all_questions = extract_questions_from_all_files()
        selected = random.sample(all_questions, min(count, len(all_questions)))
        output_path = os.path.join(GENERATED_FOLDER, "output.docx")
        generate_docx(selected, output_path)
        return send_file(output_path, as_attachment=True)
    return render_template("user_request.html")

if __name__ == '__main__':
    app.run(debug=True)
